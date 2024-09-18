from werkzeug.security import generate_password_hash, check_password_hash
from flask import Blueprint, request, jsonify
from datetime import datetime, timedelta
from pdf2image import convert_from_path
from transformers import pipeline
from functools import wraps
from PIL import Image
from app import mongo  # MongoDB instance
import pytesseract
import tempfile
import docx
import jwt
import os


main_bp = Blueprint('main', __name__)

# Set the Tesseract executable path
pytesseract.pytesseract.tesseract_cmd = r'D:\Kaif\Hackathon\Suprem court\Application\Tesseract\tesseract.exe'

# Initialize the NER pipeline outside of the route to avoid reloading the model each time
ner_pipeline = pipeline("token-classification", model="Sidziesama/Legal_NER_Support_Model")


def token_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = None
        if 'Authorization' in request.headers:
            token = request.headers['Authorization'].split(" ")[1]  # Extract token from 'Bearer <token>'

        if not token:
            return jsonify({'message': 'Token is missing!'}), 403

        try:
            # Decode the token using your secret key
            secret_key = "helloworld"
            data = jwt.decode(token, secret_key, algorithms=['HS256'])
            current_user = mongo.db.users.find_one({'_id': data['user_id']})
        except jwt.ExpiredSignatureError:
            return jsonify({'message': 'Token has expired!'}), 401
        except jwt.InvalidTokenError:
            return jsonify({'message': 'Invalid token!'}), 401

        return f(current_user, *args, **kwargs)

    return decorated

@main_bp.route('/')
def home():
    return "Welcome to PramanAI API!"

@main_bp.route('/api/protected', methods=['GET'])
@token_required
def protected_route(current_user):
    return jsonify({'message': 'This is a protected route!', 'user': current_user['username']})

@main_bp.route('/api/register', methods=['POST'])
def register():
    data = request.json
    username = data['username']
    password = generate_password_hash(data['password'], method='pbkdf2:sha256')
    
    # Check if user already exists
    existing_user = mongo.db.users.find_one({'username': username})
    if existing_user:
        return jsonify({'message': 'Username already exists'}), 400

    # Create new user
    mongo.db.users.insert_one({'username': username, 'password': password})
    return jsonify({'message': 'User registered successfully!'})

@main_bp.route('/api/login', methods=['POST'])
def login():
    data = request.json
    username = data['username']
    password = data['password']
    
    # Find user in MongoDB
    user = mongo.db.users.find_one({'username': username})
    
    if user and check_password_hash(user['password'], password):
        # Access the SECRET_KEY from the Flask app's config
        secret_key = "helloworld"

        # Generate JWT token valid for 1 hour
        token = jwt.encode({
            'user_id': str(user['_id']),  # Include user information (like id) in the payload
            'exp': datetime.utcnow() + timedelta(hours=1)  # Token expiry time
        }, secret_key, algorithm='HS256')

        return jsonify({
            'message': 'Login successful!',
            'token': token
        })

    return jsonify({'message': 'Invalid credentials!'}), 401

# Your OCR and NER routes remain unchanged
@main_bp.route('/api/ocr', methods=['POST'])
def ocr():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['file']
    lang = request.form.get('lang', 'eng')

    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    file_ext = file.filename.split('.')[-1].lower()

    # Temporary storage for the uploaded file
    with tempfile.NamedTemporaryFile(delete=False) as temp_file:
        file.save(temp_file.name)

        if file_ext == 'pdf':
            try:
                # Convert PDF to images
                images = convert_from_path(temp_file.name)
                text = ""
                for image in images:
                    text += pytesseract.image_to_string(image, lang=lang)
            except Exception as e:
                return jsonify({'error': str(e)}), 500

        elif file_ext == 'docx':
            try:
                doc = docx.Document(temp_file.name)
                text = "\n".join([para.text for para in doc.paragraphs])
            except Exception as e:
                return jsonify({'error': str(e)}), 500
        else:
            try:
                image = Image.open(temp_file.name)
                text = pytesseract.image_to_string(image, lang=lang)
            except Exception as e:
                return jsonify({'error': str(e)}), 500

    # Clean up temporary file
    os.remove(temp_file.name)

    return jsonify({'text': text})

@main_bp.route('/api/ner', methods=['POST'])
def ner():
    # Check if 'file' is in the request, otherwise assume raw text
    if 'file' in request.files:
        file = request.files['file']
        lang = request.form.get('lang', 'eng')

        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400

        file_ext = file.filename.split('.')[-1].lower()

        # Temporary storage for the uploaded file
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            file.save(temp_file.name)

            if file_ext == 'pdf':
                try:
                    # Convert PDF to images and extract text
                    images = convert_from_path(temp_file.name)
                    text = ""
                    for image in images:
                        text += pytesseract.image_to_string(image, lang=lang)
                except Exception as e:
                    return jsonify({'error': str(e)}), 500

            elif file_ext == 'docx':
                try:
                    # Extract text from DOCX
                    doc = docx.Document(temp_file.name)
                    text = "\n".join([para.text for para in doc.paragraphs])
                except Exception as e:
                    return jsonify({'error': str(e)}), 500
            else:
                try:
                    # Extract text from image
                    image = Image.open(temp_file.name)
                    text = pytesseract.image_to_string(image, lang=lang)
                except Exception as e:
                    return jsonify({'error': str(e)}), 500

        # Clean up temporary file
        os.remove(temp_file.name)

    elif request.content_type == 'application/json':
        # Handle raw text input if no file was uploaded
        data = request.json
        text = data.get('text')
        if not text:
            return jsonify({"error": "No text provided"}), 400
    else:
        return jsonify({"error": "No data provided"}), 400

    # Process the extracted text with the NER pipeline
    result = ner_pipeline(text)

    # Convert the result for JSON serialization
    entities = [
        {"word": entity['word'], "entity": entity['entity'], "score": float(entity['score'])}
        for entity in result
    ]

    return jsonify({"entities": entities}), 200
